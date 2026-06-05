"""多格式文本提取器 — 支持 .txt / .md / .docx / .epub"""

import logging

logger = logging.getLogger(__name__)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """根据文件扩展名提取纯文本。

    Returns:
        提取的纯文本字符串
    Raises:
        ValueError: 不支持的格式
        UnicodeDecodeError: 编码问题
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext in ("txt", "text", "md", "markdown"):
        return _read_plain(file_bytes)

    if ext == "docx":
        return _read_docx(file_bytes)

    if ext == "epub":
        return _read_epub(file_bytes)

    raise ValueError(f"不支持的文件格式: .{ext}，支持 .txt / .md / .docx / .epub")


def _read_plain(data: bytes) -> str:
    """读取纯文本，自动检测编码"""
    for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    raise UnicodeDecodeError("utf-8", data, 0, 1, "无法解码文本")


def _read_docx(data: bytes) -> str:
    """从 .docx 提取文本"""
    from io import BytesIO
    from docx import Document

    doc = Document(BytesIO(data))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)

    result = "\n\n".join(paragraphs)
    logger.info(f"[FormatReader] .docx → {len(result)} 字符, {len(paragraphs)} 段落")
    return result


def _read_epub(data: bytes) -> str:
    """从 .epub 提取文本"""
    from io import BytesIO
    import ebooklib
    from ebooklib import epub

    book = epub.read_epub(BytesIO(data))
    paragraphs = []

    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        try:
            content = item.get_content().decode("utf-8")
        except UnicodeDecodeError:
            try:
                content = item.get_content().decode("gbk")
            except UnicodeDecodeError:
                continue

        # 简单 HTML 标签剥离
        import re
        # 移除脚本和样式
        content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL)
        content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL)
        # 移除 HTML 标签
        content = re.sub(r'<[^>]+>', '\n', content)
        # 解码 HTML 实体
        content = content.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>')
        content = content.replace('&amp;', '&').replace('&quot;', '"').replace('&#39;', "'")
        # 清理多余空行
        content = re.sub(r'\n{3,}', '\n\n', content).strip()
        if content:
            paragraphs.append(content)

    result = "\n\n".join(paragraphs)
    logger.info(f"[FormatReader] .epub → {len(result)} 字符, {len(paragraphs)} 章节")
    return result
